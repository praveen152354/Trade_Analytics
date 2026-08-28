"""
Builds a standalone Word document summarizing the Snowflake objects
provisioned by Terraform and the scripts/files that make up the pipeline.
Run once, locally, to produce docs/Trade_Analytics_Summary.docx:

    pip install python-docx
    python docs/generate_summary_doc.py
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GREY = RGBColor(0x55, 0x55, 0x55)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = GREY


def add_h2(doc, text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = NAVY


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p_format = p.paragraph_format
    p_format.left_indent = Inches(0.25)


def build():
    doc = Document()

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    add_title(doc, "Trade Analytics — Provisioned Objects & Scripts")
    add_subtitle(doc, "Data Engineering Case Study — Snowflake + dbt + Airflow + Terraform")
    doc.add_paragraph()

    # ---------------------------------------------------------------
    add_h2(doc, "1. Snowflake objects created (via Terraform + one SQL script)")
    doc.add_paragraph(
        "All objects below were provisioned by terraform apply against the live "
        "trial account (database TRADE_ANALYTICS), organized as a medallion "
        "architecture: BRONZE (raw landing) -> SILVER (staging + business-rule "
        "evaluation) -> GOLD (marts + Type 2 SCD snapshot). One object, "
        "BRONZE.GENERATE_TRADE_FILES, is deployed by a plain SQL script "
        "(terraform/sql/generate_trade_files_procedure.sql) rather than a "
        "snowflake_procedure_python resource, keeping its Python body under direct "
        "version control in one reviewable file. Everything else is fully "
        "Terraform-managed."
    )

    add_table(
        doc,
        ["Object", "Type", "Schema", "Purpose"],
        [
            ["TRADE_ANALYTICS_WH", "Warehouse", "—", "XSMALL, auto-suspend 60s. Compute for ingestion and dbt."],
            ["TRADE_ANALYTICS", "Database", "—", "Top-level container for the whole pipeline."],
            ["BRONZE", "Schema", "TRADE_ANALYTICS", "Landing zone: raw trade files, stage, stream, tasks."],
            ["SILVER", "Schema", "TRADE_ANALYTICS", "dbt staging + business-rule evaluation (stg_trades, int_trades_evaluated)."],
            ["GOLD", "Schema", "TRADE_ANALYTICS", "dbt star schema: dim_* tables, fct_valid_trades, fct_rejected_trades, fct_trade_status, rpt_trade_report + Type 2 SCD snapshot."],
            ["TRADE_ANALYTICS_LOADER", "Role", "—", "Used by the standalone Python dev/test script (PUT + COPY INTO)."],
            ["TRADE_ANALYTICS_TRANSFORMER", "Role", "—", "Used by dbt (via dbt Cloud) to build every model."],
            ["TRADE_JSON_FORMAT", "File format", "BRONZE", "JSON, one object per line, used by COPY INTO."],
            ["TRADES_STAGE", "Internal stage", "BRONZE", "Trade batch files land here (genuine cloud object storage, Snowflake-managed)."],
            ["TRADES_RAW", "Table", "BRONZE", "Insert-only landing table (raw_payload VARIANT, file_name, loaded_at). Permanent, cluster_by=to_date(loaded_at)."],
            ["TRADES_RAW_STREAM", "Stream", "BRONZE", "CDC stream on TRADES_RAW; consumed by dbt's stg_trades model."],
            ["GENERATE_TRADE_FILES", "Procedure (Snowpark Python)", "BRONZE", "Generates a mock trade batch, writes it as .jsonl straight to the stage."],
            ["GENERATE_TRADE_FILES_TASK", "Task", "BRONZE", "Calls the procedure every 2 min (configurable). Auto-retries 2x on failure."],
            ["INGEST_TRADES_TASK", "Task", "BRONZE", "COPY INTOs new stage files every 5 min (configurable). Independent schedule from generation."],
            ["TRADE_PIPELINE_ALERT", "Email notification integration", "—", "Lets Snowflake send email for the alert below."],
            ["TASK_FAILURE_ALERT", "Alert", "BRONZE", "Every 15 min, emails if any of the three tasks failed in TASK_HISTORY."],
            ["FX_RATES_S3_INTEGRATION", "Storage integration", "—", "IAM role trust to S3 -- no long-lived AWS key stored in Snowflake."],
            ["FX_RATES_STAGE", "External stage (S3)", "BRONZE", "Points at s3://<bucket>/fx_rates/ -- manually-uploaded daily rate CSVs land here."],
            ["FX_RATES_CSV_FORMAT", "File format", "BRONZE", "CSV, skip_header=1."],
            ["FX_RATES_RAW", "Table", "BRONZE", "Append-only landing table for daily FX rates (as_of_date, currency, rate_to_usd)."],
            ["INGEST_FX_RATES_TASK", "Task", "BRONZE", "COPY INTOs new S3 files once daily (configurable). No PURGE -- IAM policy is read-only and it's the user's own bucket."],
            ["DASHBOARD_STAGE", "Stage", "GOLD", "Holds the Streamlit app file(s) -- content PUT here, same pattern as GENERATE_TRADE_FILES."],
            ["TRADE_ANALYTICS_DASHBOARD", "Streamlit app (Streamlit in Snowflake)", "GOLD", "The trade dashboard, running natively in Snowflake on TRADE_ANALYTICS_WH -- no local process. Viewable in Snowsight (Projects -> Streamlit)."],
            ["TRADE_ANALYTICS_ANALYST", "Role", "—", "Read-only, masked: SELECT on fct_trade_status/rpt_trade_report only (dbt-managed grant)."],
            ["TRADE_ANALYTICS_COMPLIANCE", "Role", "—", "Read-only, unmasked: SELECT on all of GOLD including fct_rejected_trades (dbt-managed grant)."],
            ["MASK_NOTIONAL / MASK_COUNTERPARTY", "Masking policy", "GOLD", "dbt-created and dbt-attached (on-run-start hook + post_hook); masks NOTIONAL/NOTIONAL_USD/COUNTERPARTY for any role but TRANSFORMER/COMPLIANCE/ACCOUNTADMIN."],
        ],
        col_widths=[1.7, 1.5, 1.0, 3.2],
    )

    doc.add_paragraph(
        "AWS side (terraform/aws.tf): an S3 bucket (versioned, encrypted, "
        "public access blocked), a read-only IAM policy scoped to the "
        "fx_rates/ prefix, and the IAM role Snowflake assumes. The role's "
        "trust policy needs values Snowflake only generates after the "
        "storage integration exists, and the integration needs the role's "
        "ARN as input -- resolved in one terraform apply by predicting the "
        "role's ARN from account_id + role name (deterministic) rather "
        "than the usual two-pass setup."
    )

    doc.add_paragraph(
        "Both roles are granted to the account user PRAVEENMS91 and given USAGE on "
        "the warehouse/database. The transformer role additionally gets "
        "CREATE TABLE / CREATE VIEW on SILVER and GOLD, and "
        "SELECT on all present-and-future tables/streams in BRONZE. The loader role "
        "is scoped narrowly to INSERT/SELECT on TRADES_RAW and READ/WRITE on "
        "TRADES_STAGE only. A Task's own ERROR_INTEGRATION property only accepts "
        "cloud-messaging integrations (SNS/Pub-Sub/Event Grid), not EMAIL, which "
        "is why failure alerting goes through a separate ALERT object instead."
    )

    doc.add_page_break()

    # ---------------------------------------------------------------
    add_h2(doc, "2. dbt models — business rules")
    add_table(
        doc,
        ["Model", "Materialization", "What it does"],
        [
            ["base_trades_raw (BRONZE)", "view", "Thin passthrough over TRADES_RAW -- documentation/lineage only, not the actual trades data path."],
            ["base_fx_rates_raw (BRONZE)", "view", "Thin passthrough over FX_RATES_RAW -- this one IS the real source silver.fx_rates builds from."],
            ["stg_trades (SILVER)", "incremental (append), transient", "Flattens raw JSON from the stream into typed columns."],
            ["int_trades_evaluated (SILVER)", "incremental (append), transient", "Single decision point: accepts or rejects each trade message and records why."],
            ["fx_rates (SILVER)", "incremental (merge on as_of_date+currency), transient", "Deduplicates FX_RATES_RAW to one row per date+currency -- the latest loaded value."],
            ["dim_trader / dim_book / dim_counterparty / dim_product / dim_currency (GOLD)", "table, full-refresh, transient", "One row per distinct value ever seen in int_trades_evaluated (accepted or rejected). Surrogate key via dbt_utils.generate_surrogate_key()."],
            ["dim_date (GOLD)", "table, full-refresh, transient", "Standard Kimball date dimension via dbt_utils.date_spine() (2024-01-01 to 2031-12-31)."],
            ["fct_valid_trades (GOLD)", "incremental (merge on trade_id), transient, cluster_by=maturity_date", "One row per trade_id -- the latest accepted version -- joined out to every dim_* table for surrogate-key FKs."],
            ["fct_rejected_trades (GOLD)", "incremental (append), PERMANENT, cluster_by=maturity_date", "Compliance audit log of every rejected message + reason, with the same dimensional FKs."],
            ["fct_trade_status (GOLD)", "view", "fct_valid_trades + computed ACTIVE/EXPIRED + notional_usd (convert_to_usd macro)."],
            ["rpt_trade_report (GOLD)", "view", "Flat reporting layer: fct_trade_status pre-joined to dim_date (trade + maturity date). What the dashboard queries -- no joins needed."],
            ["valid_trades_snapshot (GOLD)", "snapshot, check strategy, PERMANENT", "Type 2 SCD history of fct_valid_trades. Run monthly (1st, 01:00 UTC), not on the hourly job."],
        ],
        col_widths=[1.9, 2.3, 2.8],
    )

    doc.add_paragraph(
        "Why fx_rates exists: Snowflake's COPY INTO treats an edited file "
        "re-uploaded under the same name as new rather than overwriting, so "
        "a corrected rate file can legitimately create a second row for the "
        "same (as_of_date, currency) in FX_RATES_RAW. fx_rates resolves it "
        "with the same merge-on-latest pattern as fct_valid_trades -- verified "
        "live by deliberately re-uploading an edited rate file and "
        "confirming BRONZE kept both rows (for audit) while SILVER showed "
        "only the latest."
    )

    add_h2(doc, "2a. The GOLD star schema, and why some tables are permanent")
    doc.add_paragraph(
        "fct_valid_trades and fct_rejected_trades carry a surrogate-key "
        "foreign column for every dim_* table (trader_key, book_key, "
        "counterparty_key, product_key, currency_key, trade_date_key, "
        "maturity_date_key). gold.yml tests every one of those FKs with a "
        "relationships test against its dimension's key -- 68 tests total, "
        "up from 19 before this redesign. The natural-key text columns "
        "(trader, book, counterparty, product_type, currency) are kept on "
        "the facts too, alongside the *_key columns -- a deliberate "
        "denormalization for a handful of small, fixed-vocabulary codes, "
        "so a simple query doesn't have to join every dimension just to "
        "read them."
    )
    doc.add_paragraph(
        "dbt-snowflake defaults every table/incremental model to TRANSIENT "
        "(no Fail-safe storage) unless told otherwise. Most models here "
        "leave that default in place -- explicit now, via transient=true -- "
        "because they're a pure function of BRONZE (permanent) plus the "
        "current dbt SQL: if lost, dbt run regenerates them exactly, so "
        "Fail-safe's extra 7-day recovery window would be cost with no "
        "benefit. fct_rejected_trades and valid_trades_snapshot are the "
        "deliberate exception (transient=false): both are point-in-time "
        "compliance records of what business logic decided when a message "
        "or a snapshot ran -- if the rules ever change later, replaying "
        "BRONZE would apply new rules to old messages and silently rewrite "
        "audit history, so these two earn the extra protection a real "
        "audit record deserves."
    )
    doc.add_paragraph(
        "fct_valid_trades, fct_rejected_trades, and BRONZE.TRADES_RAW carry "
        "explicit clustering keys (cluster_by) on maturity_date / "
        "to_date(loaded_at) respectively -- illustrative at this project's "
        "row count, but aimed at the exact query pattern rpt_trade_report's "
        "dashboard filter uses (a maturity-date range). "
        "snowflake_sql/observability_toolkit.sql checks clustering health "
        "via SYSTEM$CLUSTERING_INFORMATION() and demonstrates a TEMPORARY "
        "table for session-scoped ad-hoc debugging that needs no manual "
        "cleanup."
    )

    add_table(
        doc,
        ["Rule", "Outcome"],
        [
            ["Lower version than existing", "REJECTED — reason STALE_VERSION_LOWER_THAN_EXISTING"],
            ["Same version as existing", "ACCEPTED — merges in, replacing the row"],
            ["Maturity date already in the past", "REJECTED — reason MATURITY_DATE_IN_PAST"],
            ["Maturity date passes after acceptance", "fct_trade_status view marks it EXPIRED (computed, not mutated)"],
            ["Two versions of one trade in the same batch", "Highest version ACCEPTED; the other REJECTED as SUPERSEDED_IN_BATCH"],
            ["Any rejection", "Logged to fct_rejected_trades — append-only audit trail"],
        ],
        col_widths=[3.5, 3.5],
    )

    add_h2(doc, "2b. RBAC & data masking")
    doc.add_paragraph(
        "Two read-only consumer roles sit alongside the two service roles "
        "(LOADER, TRANSFORMER): TRADE_ANALYTICS_COMPLIANCE sees everything "
        "in GOLD unmasked, including fct_rejected_trades (the audit log) "
        "and the raw fact/dimension tables. TRADE_ANALYTICS_ANALYST sees "
        "only fct_trade_status and rpt_trade_report -- not the underlying "
        "tables -- with COUNTERPARTY pseudonymized and "
        "NOTIONAL/NOTIONAL_USD rounded to the nearest $1M."
    )
    doc.add_paragraph(
        "The roles themselves are Terraform-managed (terraform/rbac.tf), "
        "but who can select what is managed by dbt: a +grants config in "
        "dbt_project.yml grants every GOLD object to COMPLIANCE by "
        "default, and the two masked models override that individually to "
        "add ANALYST -- deliberately not granted at the folder level, so "
        "it can't reach fct_valid_trades directly and read "
        "NOTIONAL/COUNTERPARTY unmasked. dbt re-applies these grants on "
        "every run via its native grants model config."
    )
    doc.add_paragraph(
        "Masking is two Snowflake masking policies "
        "(macros/create_masking_policies.sql, created idempotently via an "
        "on-run-start hook), each a CASE WHEN CURRENT_ROLE() IN (...) "
        "expression, attached via a post_hook on both consumer-facing "
        "views. Verified live by connecting as each role: rpt_trade_report "
        "returns masked values for ANALYST and real ones for COMPLIANCE, "
        "and ANALYST is denied on fct_valid_trades directly -- once "
        "secondary roles are turned off (Snowflake sessions default to "
        "secondary_roles=ALL, which combines every role granted to a "
        "user; this trial account's single user holds every role, so "
        "isolating ANALYST's own privileges for testing needs USE "
        "SECONDARY ROLES NONE first)."
    )

    doc.add_page_break()

    # ---------------------------------------------------------------
    add_h2(doc, "3. Repository layout & scripts")
    add_table(
        doc,
        ["Path", "Purpose"],
        [
            ["data_generator/generate_trades.py + load_to_snowflake.py", "Standalone dev/test path (PUT + COPY INTO from a local machine) -- not part of the scheduled production flow."],
            ["data_generator/generate_fx_rates.py", "Generates one CSV per day (realistic day-to-day rate drift, not static) for manual upload to S3 -- not an automated feed."],
            ["terraform/*.tf", "IaC for every Snowflake and AWS object listed in section 1, including the Tasks and Alert."],
            ["terraform/backend.tf + bootstrap_state_backend.py", "S3 remote state (shared between local applies and CI) + the one-time script that creates the bucket Terraform can't create for itself."],
            ["terraform/sql/generate_trade_files_procedure.sql", "The one Snowflake object deployed outside Terraform -- see section 1."],
            ["terraform/rbac.tf", "The two consumer roles (ANALYST, COMPLIANCE) + their schema/warehouse/database grants -- see section 2b."],
            ["dbt/trade_analytics/macros/create_masking_policies.sql", "Creates the two masking policies (idempotent, on-run-start hook) -- see section 2b."],
            ["dbt/trade_analytics/models/bronze/sources.yml", "Declares BRONZE as a dbt source -- Terraform + the Snowpark procedure own the actual objects."],
            ["dbt/trade_analytics/models/bronze/base_trades_raw.sql, base_fx_rates_raw.sql", "Thin passthrough views -- see section 2."],
            ["dbt/trade_analytics/models/silver/stg_trades.sql", "Consumes the BRONZE stream, flattens VARIANT to columns."],
            ["dbt/trade_analytics/models/silver/int_trades_evaluated.sql", "All business-rule logic (accept/reject + reason)."],
            ["dbt/trade_analytics/models/silver/fx_rates.sql", "Merge-dedup of FX_RATES_RAW -- see section 2."],
            ["dbt/trade_analytics/macros/convert_to_usd.sql", "Currency-conversion macro: loops over an FX-rate var to build a CASE expression."],
            ["dbt/trade_analytics/snapshots/valid_trades_snapshot.sql", "Type 2 SCD snapshot of the trade book, run monthly. Permanent table."],
            ["dbt/trade_analytics/models/gold/dim_trader.sql, dim_book.sql, dim_counterparty.sql, dim_product.sql, dim_currency.sql, dim_date.sql", "The six GOLD dimensions -- see section 2a."],
            ["dbt/trade_analytics/models/gold/fct_valid_trades.sql", "Merge-incremental table of current accepted trades, FK'd to every dimension."],
            ["dbt/trade_analytics/models/gold/fct_rejected_trades.sql", "Append-only rejected-trade audit log, FK'd to every dimension. Permanent table."],
            ["dbt/trade_analytics/models/gold/fct_trade_status.sql", "View computing ACTIVE/EXPIRED status + notional_usd. Masked for ANALYST -- see section 2b."],
            ["dbt/trade_analytics/models/gold/rpt_trade_report.sql", "Flat reporting view -- what dashboard/Summary.py queries. Masked for ANALYST -- see section 2b."],
            ["dbt Cloud (external, not a repo file)", "Hourly job: dbt run then dbt test. Separate monthly job: dbt snapshot. Reads/writes this repo via a deploy key (read-write, so the dbt Cloud IDE can branch/commit/push); separate Development and Production environments."],
            ["orchestration/airflow/ (alternative)", "Complete Docker Compose Airflow stack, documented but not the primary orchestration path."],
            ["dashboard/Summary.py + pages/1_Trade_Details.py + common.py + environment.yml", "Multi-page Streamlit report over rpt_trade_report: Summary (KPIs + small charts with dynamic captions) and Trade Details (full table, per-trade Type 2 SCD history, rejected-trades audit). Runs natively in Snowflake (Streamlit in Snowflake, terraform/streamlit.tf) -- no local process; works locally too (streamlit run) via a Snowpark-session fallback, same code either way."],
            ["snowflake_sql/observability_toolkit.sql", "Ready-to-run debugging, time-travel, optimization and monitoring queries."],
            [".github/workflows/dbt_ci.yml, terraform_ci.yml", "CI/CD: dbt build/test on PR + merge; terraform fmt/validate/plan on PR, apply on manual dispatch (sequenced after plan to avoid an S3 state-lock race -- both verified live)."],
            ["docs/SETUP_GUIDE.md, VALIDATION_LOGIC.md, SCALABILITY.md", "Step-by-step setup, rule-by-rule rationale, and the failure-handling / monitoring / 10,000x-scale write-up."],
        ],
        col_widths=[3.6, 3.4],
    )

    doc.add_page_break()

    # ---------------------------------------------------------------
    add_h2(doc, "4. Pipeline flow")
    add_code_block(
        doc,
        "BRONZE                    SILVER                GOLD\n"
        "\n"
        "GENERATE_TRADE_FILES_TASK (every 2 min)\n"
        "      |  CALL BRONZE.GENERATE_TRADE_FILES(...)\n"
        "      v\n"
        "TRADES_STAGE  (files)\n"
        "      ^\n"
        "      |  polls\n"
        "INGEST_TRADES_TASK (every 5 min) --(COPY INTO)-->  TRADES_RAW\n"
        "                                                        |\n"
        "                                                  TRADES_RAW_STREAM\n"
        "                                                        |\n"
        "                                          base_trades_raw (view, lineage only)\n"
        "                                                        |\n"
        "                                            stg_trades  (dbt Cloud,\n"
        "                                                  |      hourly job)\n"
        "                                        int_trades_evaluated\n"
        "                                              /            \\\n"
        "                              fct_valid_trades          fct_rejected_trades\n"
        "                              (FK'd to dim_trader,       (same FKs, PERMANENT,\n"
        "                               dim_book, dim_counter-     compliance audit log)\n"
        "                               party, dim_product,\n"
        "                               dim_currency, dim_date)\n"
        "                                    /        \\\n"
        "                        fct_trade_status   valid_trades_snapshot\n"
        "                        (view, incl.        (Type 2 SCD, PERMANENT,\n"
        "                         notional_usd)        dbt Cloud monthly job)\n"
        "                              |\n"
        "                       rpt_trade_report  (flat, pre-joined --\n"
        "                              |            what the dashboard queries)\n"
        "                              v\n"
        "                    Streamlit dashboard (Streamlit in Snowflake --\n"
        "                    filters: trader, book, counterparty, product,\n"
        "                    currency, status, maturity range)\n"
        "\n"
        "Manual upload --> S3 (fx_rates/) --> FX_RATES_STAGE\n"
        "                                          |  polls, once daily\n"
        "                                          v\n"
        "                              INGEST_FX_RATES_TASK --(COPY INTO)--> FX_RATES_RAW\n"
        "                                                                        |\n"
        "                                                          base_fx_rates_raw (view)\n"
        "                                                                        |\n"
        "                                                  fx_rates  (merge-dedup, dbt Cloud hourly job)",
    )

    add_h2(doc, "5. Connection details for this account")
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Account URL", "<your-org>-<your-account>.snowflakecomputing.com"],
            ["Organization / Account name", "set in your local .env (SNOWFLAKE_ORGANIZATION_NAME / SNOWFLAKE_ACCOUNT_NAME)"],
            ["Database", "TRADE_ANALYTICS"],
            ["Warehouse", "TRADE_ANALYTICS_WH"],
            ["Loader role", "TRADE_ANALYTICS_LOADER"],
            ["Transformer role (dbt)", "TRADE_ANALYTICS_TRANSFORMER"],
        ],
        col_widths=[3, 4],
    )
    doc.add_paragraph(
        "Account identifiers and credentials are never stored in this document "
        "or in the repository — they live only in the local, git-ignored .env "
        "file. See docs/SETUP_GUIDE.md for how to populate it."
    )

    doc.save("Trade_Analytics_Summary.docx")
    print("Wrote Trade_Analytics_Summary.docx")


if __name__ == "__main__":
    build()
